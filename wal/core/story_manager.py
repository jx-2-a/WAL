"""故事总控 — SQLite 版

创建、管理、导出故事。内部使用 SQLite 存储。
支持 部(Part) → 卷(Volume) → 章(Chapter) 三级结构。
"""

from datetime import datetime
from pathlib import Path
from typing import Optional

from ..models.story import Story, Chapter, Scene, StoryStatus, Volume, Part
from ..storage.database import Database
from ..storage.story_repo import StoryRepository


class StoryManager:
    """故事生命周期管理（SQLite 后端）"""

    def __init__(self, project_dir: str):
        self.project_dir = Path(project_dir)
        self.project_dir.mkdir(parents=True, exist_ok=True)

        db_path = self.project_dir / "wal.db"
        self.db = Database(str(db_path))
        self.repo = StoryRepository(self.db)

        # 自动初始化架构 + 迁移（init_schema 幂等，可安全重复调用）
        if not self.db.schema_exists():
            self.db.init_schema()
            # 尝试从 YAML 自动迁移
            if (self.project_dir / "story.yaml").exists():
                self.db.migrate_from_yaml(str(self.project_dir))
        else:
            # 已有数据库：仍然调用 init_schema 以运行增量迁移
            self.db.init_schema()

        self._story: Optional[Story] = None

    # ═══ 故事 ═════════════════════════════════════════════════════

    def create_story(self, name: str, author: str, summary: str, genre: str = "",
                     tags: list[str] | None = None) -> Story:
        """创建新故事"""
        now = datetime.now().isoformat()
        story = Story(
            name=name, author=author, summary=summary, genre=genre,
            tags=tags or [], status=StoryStatus.PLANNING,
            created_at=now, updated_at=now,
        )
        self.repo.save_story(story.model_dump(mode="json"))
        self._story = story
        return story

    def load_story(self) -> Optional[Story]:
        """从 SQLite 加载故事（含部/卷/章/场景完整层级）"""
        story_data = self.repo.load_story()
        if not story_data:
            self._story = None
            return None

        # 加载部
        parts_data = self.repo.list_parts()
        parts = []
        for part_row in parts_data:
            # 加载部下的卷
            volumes_data = self.repo.list_volumes(part_id=part_row["id"])
            volumes = []
            for vol_row in volumes_data:
                # 加载卷下的章节
                chapters_data = self.repo.list_chapters(volume_id=vol_row["id"])
                chapters = []
                for ch_row in chapters_data:
                    scenes_data = self.repo.list_scenes_by_chapter(ch_row["id"])
                    scenes = [Scene(**sc) for sc in scenes_data]
                    ch_dict = self._row_to_chapter_dict(ch_row, scenes)
                    chapters.append(Chapter(**ch_dict))
                vol = Volume(
                    id=vol_row["id"], part_id=vol_row.get("part_id", ""),
                    number=vol_row["number"], title=vol_row["title"],
                    summary=vol_row["summary"], theme=vol_row.get("theme", ""),
                    status=vol_row.get("status", "planning"),
                    notes=vol_row.get("notes", ""),
                    chapters=chapters,
                )
                volumes.append(vol)
            part = Part(
                id=part_row["id"], number=part_row["number"],
                title=part_row["title"], summary=part_row.get("summary", ""),
                notes=part_row.get("notes", ""),
                volumes=volumes,
            )
            parts.append(part)

        # 查找不属于任何卷的章节（向后兼容：直接挂在 story 下）
        all_vol_chapter_ids = set()
        for part in parts:
            for vol in part.volumes:
                for ch in vol.chapters:
                    all_vol_chapter_ids.add(ch.id)

        orphan_chapters_data = self.repo.list_chapters()  # 所有章节
        orphan_chapters = []
        for ch_row in orphan_chapters_data:
            if ch_row["id"] not in all_vol_chapter_ids:
                scenes_data = self.repo.list_scenes_by_chapter(ch_row["id"])
                scenes = [Scene(**sc) for sc in scenes_data]
                ch_dict = self._row_to_chapter_dict(ch_row, scenes)
                orphan_chapters.append(Chapter(**ch_dict))

        story_data["parts"] = parts
        story_data["chapters"] = orphan_chapters
        self._story = Story(**story_data)
        return self._story

    def _row_to_chapter_dict(self, ch_row: dict, scenes: list[Scene]) -> dict:
        """将 DB 行 + 场景列表转为 Chapter 构造所需的 dict"""
        return {
            "id": ch_row["id"],
            "volume_id": ch_row.get("volume_id") or "",
            "number": ch_row["number"],
            "title": ch_row["title"],
            "status": ch_row["status"],
            "summary": ch_row["summary"],
            "word_count_target": ch_row["word_count_target"],
            "actual_word_count": ch_row["actual_word_count"],
            "plot_points_involved": ch_row["plot_points_involved"],
            "character_appearances": ch_row["character_appearances"],
            "notes": ch_row["notes"],
            "scenes": scenes,
        }

    def get_story(self) -> Optional[Story]:
        """获取当前故事（自动加载）"""
        if self._story is None:
            self.load_story()
        return self._story

    def update_story(self, **kwargs) -> Story:
        """更新故事元数据"""
        if not self._story:
            raise ValueError("No story loaded")
        for key, value in kwargs.items():
            if hasattr(self._story, key):
                setattr(self._story, key, value)
        self._story.updated_at = datetime.now().isoformat()
        self.repo.save_story(self._story.model_dump(mode="json"))
        return self._story

    # ═══ 部/篇 ═════════════════════════════════════════════════════

    def _ensure_story_exists(self) -> None:
        """确保 stories 表有主记录（FK 约束需要）"""
        if not self.repo.story_exists():
            now = datetime.now().isoformat()
            self.repo.save_story({
                "id": "main", "name": "", "author": "",
                "status": "planning", "created_at": now, "updated_at": now,
            })

    def add_part(self, title: str, summary: str = "", notes: str = "") -> Part:
        """添加新部/篇"""
        self._ensure_story_exists()
        next_num = self.repo.next_part_number()
        pid = f"part_{next_num:03d}"
        part = Part(id=pid, number=next_num, title=title,
                    summary=summary, notes=notes)
        self.repo.save_part(part.model_dump(mode="json"))
        if self._story:
            self._story.parts.append(part)
        return part

    def get_part(self, part_id: str) -> Optional[Part]:
        """获取指定部"""
        if self._story:
            for p in self._story.parts:
                if p.id == part_id:
                    return p
        part_row = self.repo.load_part(part_id)
        if not part_row:
            return None
        return Part(**part_row)

    def list_parts(self) -> list[Part]:
        """列出所有部"""
        self.load_story()
        return self._story.parts if self._story else []

    # ═══ 卷 ═════════════════════════════════════════════════════

    def add_volume(self, title: str, part_id: str = "", summary: str = "",
                   theme: str = "", notes: str = "") -> Volume:
        """添加新卷。可指定所属部。

        Args:
            title: 卷标题
            part_id: 所属部ID（可选）
            summary: 卷摘要
            theme: 卷主题
            notes: 备注
        """
        self._ensure_story_exists()
        next_num = self.repo.next_volume_number(part_id)
        vid = f"vol_{next_num:03d}"
        volume = Volume(
            id=vid, part_id=part_id, number=next_num,
            title=title, summary=summary, theme=theme,
            notes=notes,
        )
        self.repo.save_volume(volume.model_dump(mode="json"))
        if self._story:
            if part_id:
                for part in self._story.parts:
                    if part.id == part_id:
                        part.volumes.append(volume)
                        break
        return volume

    def delete_volume(self, volume_id: str) -> dict:
        """删除指定卷及其所有章节和场景

        同时清理 FTS 全文索引和内存模型。
        如果卷下有章节，会一并删除所有章节及其场景。

        Args:
            volume_id: 卷ID（如 vol_001）

        Returns:
            {"deleted": True, "volume_id": "vol_001"}
            或 {"error": "..."}
        """
        self._ensure_story_exists()
        vol = self.get_volume(volume_id)
        if not vol:
            return {"error": f"卷 {volume_id} 不存在"}
        # 删除卷下所有章节
        if self._story:
            for ch in list(self._story.chapters):
                if getattr(ch, 'volume_id', '') == volume_id:
                    self.delete_chapter(ch.number)
        # 从内存模型中移除
        if self._story:
            for part in self._story.parts:
                part.volumes = [v for v in part.volumes if v.id != volume_id]
        # 从数据库删除
        self.repo.delete_volume(volume_id)
        return {"deleted": True, "volume_id": volume_id, "title": vol.title}

    def get_volume(self, volume_id: str) -> Optional[Volume]:
        """获取指定卷"""
        if self._story:
            for part in self._story.parts:
                for v in part.volumes:
                    if v.id == volume_id:
                        return v
        vol_row = self.repo.load_volume(volume_id)
        if not vol_row:
            return None
        return self._row_to_volume(vol_row)

    def update_volume(self, volume_id: str, **kwargs) -> Volume:
        """更新卷信息（title, summary, theme, status, notes）"""
        vol = self.get_volume(volume_id)
        if not vol:
            raise ValueError(f"Volume '{volume_id}' not found")
        for key, value in kwargs.items():
            if hasattr(vol, key):
                setattr(vol, key, value)
                self.repo.update_volume_field(volume_id, key, value)
        # 同步更新内存模型
        if self._story:
            for part in self._story.parts:
                for i, v in enumerate(part.volumes):
                    if v.id == volume_id:
                        part.volumes[i] = vol
                        break
        return vol

    def list_volumes(self, part_id: str = "") -> list[Volume]:
        """列出卷。可按部过滤。包含各卷的章节列表。"""
        # 优先从内存返回
        if self._story:
            candidates = []
            if not part_id:
                for part in self._story.parts:
                    candidates.extend(part.volumes)
            else:
                for part in self._story.parts:
                    if part.id == part_id:
                        candidates = part.volumes
                        break
            if candidates and all(len(v.chapters) > 0 for v in candidates if v.id):
                return candidates
            # 部分卷可能有章节，尝试补充加载
            if candidates:
                for v in candidates:
                    if not v.chapters:
                        ch_rows = self.repo.list_chapters(volume_id=v.id)
                        v.chapters = [Chapter(**self._row_to_chapter_dict(
                            ch_row, [Scene(**sc) for sc in
                                     self.repo.list_scenes_by_chapter(ch_row["id"])]
                        )) for ch_row in ch_rows]
                return candidates

        # 从 DB 加载
        vol_rows = self.repo.list_volumes(part_id)
        result = []
        for v_row in vol_rows:
            vol = self._row_to_volume(v_row)
            ch_rows = self.repo.list_chapters(volume_id=v_row["id"])
            vol.chapters = [Chapter(**self._row_to_chapter_dict(
                ch_row, [Scene(**sc) for sc in
                         self.repo.list_scenes_by_chapter(ch_row["id"])]
            )) for ch_row in ch_rows]
            result.append(vol)
        return result

    def _row_to_volume(self, row: dict) -> Volume:
        """将 DB 行转为 Volume 模型（处理 NULL → ''）"""
        row = dict(row)
        row["part_id"] = row.get("part_id") or ""
        return Volume(**row)

    def get_volume_context(self, volume_id: str) -> dict:
        """获取卷级写作上下文（卷摘要、主题、章节列表、进度）"""
        vol_row = self.repo.load_volume(volume_id)
        if not vol_row:
            return {"error": f"卷 {volume_id} 不存在"}

        chapters_data = self.repo.list_chapters(volume_id=volume_id)
        chapters_info = []
        total_words = 0
        for ch_row in chapters_data:
            status_label = {"draft": "草稿", "writing": "写作中", "done": "完成"}.get(
                ch_row["status"], ch_row["status"])
            chapters_info.append({
                "number": ch_row["number"],
                "title": ch_row["title"],
                "status": status_label,
                "summary": ch_row["summary"],
                "words": ch_row["actual_word_count"],
            })
            total_words += ch_row["actual_word_count"]

        done_count = sum(1 for c in chapters_info if c["status"] == "完成")

        # 查找所属部
        part_name = ""
        if vol_row.get("part_id"):
            part_row = self.repo.load_part(vol_row["part_id"])
            if part_row:
                part_name = part_row["title"]

        return {
            "volume_id": volume_id,
            "volume_title": vol_row["title"],
            "volume_number": vol_row["number"],
            "part_name": part_name,
            "theme": vol_row.get("theme", ""),
            "summary": vol_row.get("summary", ""),
            "status": vol_row.get("status", "planning"),
            "chapter_count": len(chapters_info),
            "done_chapters": done_count,
            "total_words": total_words,
            "chapters": chapters_info,
        }

    # ═══ 章节 ═════════════════════════════════════════════════════

    def add_chapter(self, title: str, word_count_target: int = 3000,
                    summary: str = "", notes: str = "",
                    volume_id: str = "", volume_number: int = 0,
                    chapter_number: int = 0) -> Chapter:
        """添加新章节。可通过 volume_id 或 volume_number 指定所属卷。

        Args:
            title: 章节标题
            word_count_target: 目标字数
            summary: 章节摘要
            notes: 备注
            volume_id: 卷ID（如 vol_001）
            volume_number: 卷序号（自动查找对应卷ID，volume_id 优先）
            chapter_number: 章节号（0=自动分配到末尾，>0=使用指定序号）
        """
        self._ensure_story_exists()
        # 如果传了 volume_number 但没有 volume_id，查找卷ID
        if not volume_id and volume_number > 0:
            vol = self.repo.get_volume_by_number(volume_number)
            if vol:
                volume_id = vol["id"]
            else:
                raise ValueError(f"卷 {volume_number} 不存在，请先创建卷。")

        if chapter_number > 0:
            # 检查指定章节号是否已存在
            existing = self.repo.load_chapter_by_number(chapter_number, volume_id)
            if existing:
                raise ValueError(f"第 {chapter_number} 章已存在，请使用其他章节号或调用 update_chapter_info 修改")
            next_num = chapter_number
        else:
            next_num = self.repo.next_chapter_number(volume_id) if volume_id else \
                       self.repo.next_chapter_number()
        ch_id = f"ch_{next_num:04d}"
        chapter = Chapter(
            id=ch_id, volume_id=volume_id, number=next_num, title=title,
            word_count_target=word_count_target,
            summary=summary, notes=notes,
        )
        ch_dict = chapter.model_dump(mode="json")
        ch_dict["id"] = ch_id
        ch_dict["volume_id"] = volume_id
        self.repo.save_chapter(ch_dict)

        # 更新内存模型
        if self._story:
            self._story.chapters.append(chapter)
        return chapter

    def get_chapter(self, number: int, volume_id: str = "",
                    volume_number: int = 0) -> Optional[Chapter]:
        """获取指定章节。可通过 volume_id 或 volume_number 限定卷范围。

        Args:
            number: 章节序号
            volume_id: 卷ID（可选，限定搜索范围）
            volume_number: 卷序号（volume_id 优先）
        """
        if not volume_id and volume_number > 0:
            vol = self.repo.get_volume_by_number(volume_number)
            if vol:
                volume_id = vol["id"]
        ch_row = self.repo.load_chapter_by_number(number, volume_id)
        if not ch_row:
            return None
        scenes_data = self.repo.list_scenes_by_chapter(ch_row["id"])
        scenes = [Scene(**sc) for sc in scenes_data]
        return Chapter(**self._row_to_chapter_dict(ch_row, scenes))

    def update_chapter(self, number: int, **kwargs) -> Chapter:
        """更新章节信息"""
        ch = self.get_chapter(number)
        if not ch:
            raise ValueError(f"Chapter {number} not found")
        ch_id = f"ch_{number:04d}"
        for key, value in kwargs.items():
            if hasattr(ch, key):
                setattr(ch, key, value)
                if key == "actual_word_count" or key not in (
                    "scenes", "plot_points_involved", "character_appearances"
                ):
                    self.repo.update_chapter_field(ch_id, key, value)
        return ch

    def set_chapter_status(self, number: int, status: str) -> Chapter:
        """设置章节状态（draft / writing / done）"""
        ch_id = f"ch_{number:04d}"
        self.repo.update_chapter_field(ch_id, "status", status)
        return self.get_chapter(number)

    def delete_chapter(self, number: int) -> dict:
        """删除指定章节及其所有场景

        同时从 FTS 索引中移除场景内容。
        **级联清理**: 同时删除该章的角色快照、情节点关联、伏笔引用。
        会同步更新内存模型（如果已加载）。
        """
        ch_id = f"ch_{number:04d}"

        # —— 级联清理：角色快照 ——
        from ..storage.char_repo import CharacterRepository
        char_repo = CharacterRepository(self.db)
        snapshots_deleted = char_repo.delete_snapshots_by_chapter(number)

        # —— 级联清理：情节点 + 伏笔引用 ——
        from ..storage.plot_repo import PlotRepository
        plot_repo = PlotRepository(self.db)
        points_deleted = plot_repo.delete_points_by_chapter(number)
        fw_reset = plot_repo.reset_foreshadowing_chapter(number)

        # —— 原有清理：FTS + 场景 + 章节 ——
        scenes = self.repo.list_scenes_by_chapter(ch_id)
        for sc in scenes:
            self.repo.remove_scene_from_fts(sc["id"])
        self.repo.delete_scenes_by_chapter(ch_id)
        deleted = self.repo.delete_chapter(ch_id)
        if deleted == 0:
            return {"error": f"第{number}章不存在"}

        # 更新内存模型
        if self._story:
            self._story.chapters = [
                ch for ch in self._story.chapters if ch.number != number
            ]

        return {
            "deleted": True,
            "chapter_number": number,
            "chapter_id": ch_id,
            "cascade_cleanup": {
                "character_snapshots": snapshots_deleted,
                "plot_points": points_deleted,
                "foreshadowing_refs_cleared": fw_reset["cleared_created"] + fw_reset["cleared_resolved"],
            },
        }

    # ═══ 场景 ═════════════════════════════════════════════════════

    def add_scene(self, chapter_number: int, title: str, location_id: str = "",
                  time_point: str = "", characters_present: list[str] | None = None,
                  notes: str = "") -> Scene:
        """向章节添加场景"""
        ch_id = f"ch_{chapter_number:04d}"
        scene_index = self.repo.next_scene_index(ch_id)
        scene = Scene(
            id=f"sc_ch{chapter_number}_{scene_index+1:02d}",
            title=title, location_id=location_id,
            time_point=time_point,
            characters_present=characters_present or [],
            notes=notes,
        )
        sc_dict = scene.model_dump(mode="json")
        sc_dict["chapter_id"] = ch_id
        sc_dict["scene_index"] = scene_index
        self.repo.save_scene(sc_dict)
        return scene

    def update_scene_content(self, chapter_number: int, scene_index: int,
                             content: str) -> Scene:
        """更新场景正文。若章节尚无场景且 scene_index=0，自动创建默认场景。"""
        ch_id = f"ch_{chapter_number:04d}"
        scenes = self.repo.list_scenes_by_chapter(ch_id)

        # 自动创建场景：章节无场景且请求 scene_index=0
        if len(scenes) == 0 and scene_index == 0:
            self.add_scene(chapter_number, title=f"场景{chapter_number}-1")
            scenes = self.repo.list_scenes_by_chapter(ch_id)

        if scene_index < 0 or scene_index >= len(scenes):
            raise ValueError(f"Scene index {scene_index} out of range (chapter has {len(scenes)} scenes)")
        scene_data = scenes[scene_index]
        scene_id = scene_data["id"]

        # 直接写入 SQLite
        word_count = len(content)
        self.repo.update_scene_content(scene_id, content)

        # 更新章节总字数
        self._recalc_chapter_word_count(ch_id)

        # 更新 FTS 索引
        ch_row = self.repo.load_chapter_by_number(chapter_number)
        if ch_row:
            self.repo.index_scene_for_fts(
                chapter_id=ch_id,
                chapter_title=ch_row.get("title", ""),
                chapter_summary=ch_row.get("summary", ""),
                scene_id=scene_id,
                scene_title=scene_data.get("title", ""),
                content=content,
                characters=", ".join(scene_data.get("characters_present", [])),
                location=scene_data.get("location_id", ""),
                plot_refs=", ".join(scene_data.get("plot_advancements", [])),
            )

        return Scene(
            id=scene_id, title=scene_data["title"],
            location_id=scene_data["location_id"],
            time_point=scene_data["time_point"],
            characters_present=scene_data["characters_present"],
            content=content, plot_advancements=scene_data["plot_advancements"],
            notes=scene_data["notes"], word_count=word_count,
        )

    def _recalc_chapter_word_count(self, chapter_id: str) -> None:
        """重新计算章节实际字数"""
        scenes = self.repo.list_scenes_by_chapter(chapter_id)
        total = sum(sc.get("word_count", 0) for sc in scenes)
        self.repo.update_chapter_field(chapter_id, "actual_word_count", total)

    # ═══ 查询 ═════════════════════════════════════════════════════

    def get_story_status(self) -> dict:
        """获取故事整体进度（含部/卷统计）"""
        stats = self.repo.get_story_stats()
        part_count = self.repo._count("parts", "story_id = ?", ("main",))
        vol_count = self.repo._count("volumes", "story_id = ?", ("main",))

        if stats["total_chapters"] == 0:
            story_data = self.repo.load_story()
            if not story_data:
                return {"status": "no story loaded"}
            return {
                "name": story_data.get("name", ""),
                "status": story_data.get("status", "planning"),
                "style": story_data.get("style", ""),
                "total_chapters": 0, "done_chapters": 0,
                "progress_percent": 0.0, "total_words": 0,
                "parts": part_count, "volumes": vol_count,
            }
        story_data = self.repo.load_story() or {}
        return {
            "name": story_data.get("name", ""),
            "status": story_data.get("status", "planning"),
            "style": story_data.get("style", ""),
            "total_chapters": stats["total_chapters"],
            "done_chapters": stats["done_chapters"],
            "progress_percent": stats["progress_percent"],
            "total_words": stats["total_words"],
            "parts": part_count, "volumes": vol_count,
        }

    def export_outline(self) -> str:
        """导出大纲文本 — 部 > 卷 > 章 三级结构"""
        story_data = self.repo.load_story()
        if not story_data:
            return ""

        lines = [
            f"# 《{story_data.get('name', '')}》大纲",
            f"作者：{story_data.get('author', '')}",
            f"类型：{story_data.get('genre', '')}",
            f"简介：{story_data.get('summary', '')}",
            "",
        ]

        # 重新加载获取完整的层级结构
        self.load_story()
        if not self._story:
            return "\n".join(lines)

        has_volumes = False
        has_parts = bool(self._story.parts)

        if has_parts:
            for part in self._story.parts:
                lines.append(f"## 第{part.number}部：{part.title}")
                if part.summary:
                    lines.append(f"> {part.summary}")
                lines.append("")
                has_volumes = has_volumes or bool(part.volumes)

                for vol in part.volumes:
                    self._format_volume_outline(vol, lines, indent="")
                lines.append("")
        else:
            # 无部结构：检查是否有独立卷
            all_vol_rows = self.repo.list_volumes()
            if all_vol_rows:
                for v_row in all_vol_rows:
                    vol = self._row_to_volume(v_row)
                    ch_rows = self.repo.list_chapters(volume_id=v_row["id"])
                    chapters = []
                    for ch_row in ch_rows:
                        scenes = self.repo.list_scenes_by_chapter(ch_row["id"])
                        chapters.append(Chapter(**self._row_to_chapter_dict(ch_row, [Scene(**s) for s in scenes])))
                    vol.chapters = chapters
                    self._format_volume_outline(vol, lines, indent="")
                    has_volumes = True

        if not has_volumes and not has_parts:
            # 最简模式：直接列章节（向后兼容）
            lines.append("## 章节概览")
            lines.append("")
            chapters = self.repo.list_chapters()
            for ch in chapters:
                self._format_chapter_outline(ch, lines, indent="### ")

        # 总统计
        stats = self.repo.get_story_stats()
        lines.append("---")
        lines.append(f"总计：{stats['total_chapters']} 章 | "
                     f"完成 {stats['done_chapters']} 章 | "
                     f"{stats['total_words']} 字 | "
                     f"进度 {stats['progress_percent']}%")

        return "\n".join(lines)

    def _format_volume_outline(self, vol: Volume, lines: list[str], indent: str) -> None:
        """格式化卷的大纲行"""
        vol_status = {"planning": "[规划中]", "writing": "[写作中]", "completed": "[完成]"}.get(
            vol.status, vol.status)
        lines.append(f"{indent}### 第{vol.number}卷：{vol.title}  {vol_status}")
        if vol.theme:
            lines.append(f"{indent}主题：{vol.theme}")
        if vol.summary:
            lines.append(f"{indent}> {vol.summary}")
        lines.append("")

        # 卷内章节
        for ch in vol.chapters:
            self._format_chapter_outline_chapter(ch, lines, indent=f"{indent}  ")

        # 如果 vol.chapters 为空，从 DB 加载
        if not vol.chapters and vol.id:
            ch_rows = self.repo.list_chapters(volume_id=vol.id)
            for ch_row in ch_rows:
                self._format_chapter_outline(ch_row, lines, indent=f"{indent}  ")

    def _format_chapter_outline(self, ch_row: dict, lines: list[str], indent: str) -> None:
        """格式化章节的大纲行（从 DB row）"""
        status_label = {"draft": "[草稿]", "writing": "[写作中]", "done": "[完成]"}.get(
            ch_row.get("status", ""), ch_row.get("status", ""))
        lines.append(f"{indent}第{ch_row['number']}章 {ch_row['title']}  {status_label}")
        if ch_row.get("summary"):
            lines.append(f"{indent}  摘要：{ch_row['summary']}")
        lines.append(f"{indent}  字数：{ch_row['actual_word_count']}/{ch_row['word_count_target']}")

        scenes = self.repo.list_scenes_by_chapter(ch_row["id"])
        for s in scenes:
            lines.append(f"{indent}    - {s['title']} [{s['time_point']}] ({s['word_count']}字)")

    def _format_chapter_outline_chapter(self, ch, lines: list[str], indent: str) -> None:
        """格式化章节的大纲行（从 Chapter 模型）"""
        status_label = {"draft": "[草稿]", "writing": "[写作中]", "done": "[完成]"}.get(
            ch.status, ch.status)
        lines.append(f"{indent}第{ch.number}章 {ch.title}  {status_label}")
        if ch.summary:
            lines.append(f"{indent}  摘要：{ch.summary}")
        lines.append(f"{indent}  字数：{ch.actual_word_count}/{ch.word_count_target}")
        for s in ch.scenes:
            lines.append(f"{indent}    - {s.title} [{s.time_point}] ({s.word_count}字)")

    # ═══ 导出 ═════════════════════════════════════════════════════

    def export_chapter_markdown(self, chapter_number: int) -> str:
        """导出单章为 Markdown"""
        ch = self.get_chapter(chapter_number)
        if not ch:
            return ""
        lines = [
            f"# 第{ch.number}章 {ch.title}",
            "",
        ]
        if ch.summary:
            lines.append(f"> *{ch.summary}*")
            lines.append("")
        # 场景内容：多场景用分割线隔开，不写"场景N"（避免出戏）
        has_content = any(sc.content.strip() for sc in ch.scenes)
        multi_scene = len([sc for sc in ch.scenes if sc.content.strip()]) > 1
        first = True
        for sc in ch.scenes:
            if not sc.content.strip():
                continue
            if not first and multi_scene:
                lines.append("")
                lines.append("——————")
                lines.append("")
            first = False
            lines.append(sc.content)
            lines.append("")
        if not has_content:
            lines.append("（本章暂无正文）")
            lines.append("")
        return "\n".join(lines)

    def export_chapter_html(self, chapter_number: int) -> str:
        """导出单章为 HTML"""
        ch = self.get_chapter(chapter_number)
        if not ch:
            return ""
        parts = [
            '<!DOCTYPE html><html><head><meta charset="utf-8">',
            f'<title>第{ch.number}章 {ch.title}</title>',
            '<style>body{max-width:800px;margin:0 auto;padding:20px;',
            'font-family:"Microsoft YaHei",sans-serif;line-height:1.8}',
            'h1{text-align:center}hr.scene-divider{width:40%;margin:30px auto;',
            'border:none;border-top:1px solid #ccc;text-align:center}',
            'hr.scene-divider::after{content:"——————";display:inline-block;',
            'position:relative;top:-0.7em;background:#fff;padding:0 10px;color:#999}',
            '</style></head><body>',
            f'<h1>第{ch.number}章 {ch.title}</h1>',
        ]
        if ch.summary:
            parts.append(f'<blockquote>{ch.summary}</blockquote>')
        # 场景内容：多场景用 <hr> 隔开，不写"场景N"
        scenes_with_content = [sc for sc in ch.scenes if sc.content.strip()]
        multi = len(scenes_with_content) > 1
        first = True
        for sc in scenes_with_content:
            if not first and multi:
                parts.append('<hr class="scene-divider">')
            first = False
            parts.append(f'<p>{sc.content.replace(chr(10), "<br>")}</p>')
        if not scenes_with_content:
            parts.append('<p><em>（本章暂无正文）</em></p>')
        parts.append('</body></html>')
        return "\n".join(parts)

    def export_chapter_plain(self, chapter_number: int) -> str:
        """导出单章为纯文本（网文标准格式）"""
        ch = self.get_chapter(chapter_number)
        if not ch:
            return ""
        parts = [
            f"第{ch.number}章 {ch.title}",
            "",
        ]
        # 场景内容：多场景用分割线隔开，不写"场景N"（避免出戏）
        scenes_with_content = [sc for sc in ch.scenes if sc.content.strip()]
        multi = len(scenes_with_content) > 1
        first = True
        for sc in scenes_with_content:
            if not first and multi:
                parts.append("")
                parts.append("——————")
                parts.append("")
            first = False
            parts.append(sc.content)
        if not scenes_with_content:
            parts.append("（本章暂无正文）")
        return "\n".join(parts)

    def batch_export(self, start_ch: int, end_ch: int, fmt: str = "markdown",
                     output_dir: str = "") -> dict[str, str]:
        """批量导出章节

        Args:
            start_ch: 起始章节号
            end_ch: 结束章节号（含）
            fmt: 导出格式 — markdown / html / plain
            output_dir: 输出目录（为空则返回内容 dict）

        Returns:
            {chapter_number: content_string} 或写入文件后返回 {chapter_number: filepath}
        """
        exporters = {
            "markdown": self.export_chapter_markdown,
            "html": self.export_chapter_html,
            "plain": self.export_chapter_plain,
        }
        exporter = exporters.get(fmt, self.export_chapter_markdown)
        ext = {"markdown": ".md", "html": ".html", "plain": ".txt"}.get(fmt, ".md")
        results = {}

        chapters = self.repo.list_chapters()
        for ch_row in chapters:
            ch_num = ch_row.get("number", 0)
            if start_ch <= ch_num <= end_ch:
                content = exporter(ch_num)
                if not content:
                    continue
                if output_dir:
                    import os
                    os.makedirs(output_dir, exist_ok=True)
                    fname = f"ch_{ch_num:04d}_{self._safe_filename(ch_row.get('title', ''))}{ext}"
                    fpath = os.path.join(output_dir, fname)
                    with open(fpath, "w", encoding="utf-8") as f:
                        f.write(content)
                    results[str(ch_num)] = fpath
                else:
                    results[str(ch_num)] = content

        return results

    def export_volume(self, volume_number: int, fmt: str = "markdown") -> str:
        """导出整卷的章节内容

        Args:
            volume_number: 卷序号
            fmt: 导出格式 — markdown / html / plain

        Returns:
            整卷合并后的内容字符串
        """
        vol_row = self.repo.get_volume_by_number(volume_number)
        if not vol_row:
            return f"（第{volume_number}卷不存在）"

        vol = self._row_to_volume(vol_row)
        chapters = self.repo.list_chapters(volume_id=vol.id)
        if not chapters:
            return f"# 第{volume_number}卷 {vol.title}\n\n（本卷暂无章节内容）"

        parts = []
        # 卷标题
        if fmt == "markdown":
            parts.append(f"# 第{vol.number}卷：{vol.title}")
            parts.append("")
            if vol.theme:
                parts.append(f"> 主题：{vol.theme}")
            parts.append("")
            if vol.summary:
                parts.append(f"> {vol.summary}")
                parts.append("")
            parts.append("---")
            parts.append("")
            for ch_row in chapters:
                parts.append(self.export_chapter_markdown(ch_row["number"]))
                parts.append("")
                parts.append("---")
                parts.append("")
        elif fmt == "html":
            parts.append('<!DOCTYPE html><html><head><meta charset="utf-8">')
            parts.append(f'<title>第{vol.number}卷 {vol.title}</title>')
            parts.append('<style>body{max-width:800px;margin:0 auto;padding:20px;')
            parts.append('font-family:"Microsoft YaHei",sans-serif;line-height:1.8}')
            parts.append('h1{text-align:center}h2{color:#333}</style></head><body>')
            parts.append(f'<h1>第{vol.number}卷：{vol.title}</h1>')
            if vol.summary:
                parts.append(f'<blockquote>{vol.summary}</blockquote>')
            parts.append('<hr>')
            for ch_row in chapters:
                ch_html = self.export_chapter_html(ch_row["number"])
                # 去掉 HTML 头部（跳过第一行 doctype+html+head 直到 <body>）
                body_start = ch_html.find("<body>")
                if body_start >= 0:
                    ch_html = ch_html[body_start + 6:]
                body_end = ch_html.rfind("</body>")
                if body_end >= 0:
                    ch_html = ch_html[:body_end]
                parts.append(ch_html)
                parts.append("<hr>")
            parts.append('</body></html>')
        else:
            parts.append(f"第{vol.number}卷：{vol.title}")
            parts.append("=" * 40)
            parts.append("")
            for ch_row in chapters:
                parts.append(self.export_chapter_plain(ch_row["number"]))
                parts.append("")
                parts.append("-" * 40)
                parts.append("")

        return "\n".join(parts)

    def export_full_novel(self, fmt: str = "markdown") -> str:
        """导出整部小说的完整内容

        Args:
            fmt: 导出格式 — markdown / html / plain

        Returns:
            全书内容的完整字符串
        """
        story = self.repo.load_story()
        title = story.get("name", "未命名") if story else "未命名"
        author = story.get("author", "") if story else ""

        parts = []
        if fmt == "markdown":
            parts.append(f"# 《{title}》")
            if author:
                parts.append(f"> 作者：{author}")
            parts.append("")
            parts.append("---")
            parts.append("")

            # 检测结构
            has_parts = len(self.repo.list_parts()) > 0
            has_volumes = len(self.repo.list_volumes()) > 0

            all_chapters = self.repo.list_chapters()
            all_chapters.sort(key=lambda c: c.get("number", 0))

            if has_parts:
                for p_row in self.repo.list_parts():
                    p_num = p_row.get("number", 0)
                    parts.append(f"# 第{p_num}部：{p_row.get('title', '')}")
                    parts.append("")
                    if p_row.get("summary"):
                        parts.append(f"> {p_row['summary']}")
                    parts.append("")

                    for v_row in self.repo.list_volumes(part_id=p_row["id"]):
                        vol = self._row_to_volume(v_row)
                        ch_rows = self.repo.list_chapters(volume_id=vol.id)
                        if ch_rows:
                            parts.append(f"## 第{vol.number}卷：{vol.title}")
                            parts.append("")
                            if vol.theme:
                                parts.append(f"> 主题：{vol.theme}")
                            parts.append("")
                            for ch_row in ch_rows:
                                parts.append(self.export_chapter_markdown(ch_row["number"]))
                                parts.append("")
                                parts.append("---")
                                parts.append("")
            elif has_volumes:
                for v_row in self.repo.list_volumes():
                    vol = self._row_to_volume(v_row)
                    ch_rows = self.repo.list_chapters(volume_id=vol.id)
                    if ch_rows:
                        parts.append(f"# 第{vol.number}卷：{vol.title}")
                        parts.append("")
                        if vol.theme:
                            parts.append(f"> 主题：{vol.theme}")
                        parts.append("")
                        for ch_row in ch_rows:
                            parts.append(self.export_chapter_markdown(ch_row["number"]))
                            parts.append("")
                            parts.append("---")
                            parts.append("")

                # 无归属章节
                vol_ch_ids = set()
                for v_row in self.repo.list_volumes():
                    for ch in self.repo.list_chapters(volume_id=v_row["id"]):
                        vol_ch_ids.add(ch["id"])
                orphan = [c for c in all_chapters if c["id"] not in vol_ch_ids]
                if orphan:
                    parts.append("# 独立章节")
                    parts.append("")
                    for ch_row in orphan:
                        parts.append(self.export_chapter_markdown(ch_row["number"]))
                        parts.append("")
                        parts.append("---")
                        parts.append("")
            else:
                # 纯章节模式
                for ch_row in all_chapters:
                    parts.append(self.export_chapter_markdown(ch_row["number"]))
                    parts.append("")
                    parts.append("---")
                    parts.append("")

        elif fmt == "html":
            parts.append('<!DOCTYPE html><html><head><meta charset="utf-8">')
            parts.append(f'<title>{title}</title>')
            parts.append('<style>body{max-width:800px;margin:0 auto;padding:20px;')
            parts.append('font-family:"Microsoft YaHei",sans-serif;line-height:1.8}')
            parts.append('h1{text-align:center}h2{color:#333}hr{margin:30px 0}')
            parts.append('.chapter{border-bottom:1px solid #eee;padding-bottom:20px}</style>')
            parts.append('</head><body>')
            parts.append(f'<h1>{title}</h1>')
            if author:
                parts.append(f'<p>作者：{author}</p>')
            parts.append('<hr>')

            all_chapters = self.repo.list_chapters()
            all_chapters.sort(key=lambda c: c.get("number", 0))
            for ch_row in all_chapters:
                ch_html = self.export_chapter_html(ch_row["number"])
                body_start = ch_html.find("<body>")
                if body_start >= 0:
                    ch_html = ch_html[body_start + 6:]
                body_end = ch_html.rfind("</body>")
                if body_end >= 0:
                    ch_html = ch_html[:body_end]
                parts.append('<div class="chapter">')
                parts.append(ch_html)
                parts.append('</div>')
            parts.append('</body></html>')

        else:
            parts.append(f"《{title}》")
            if author:
                parts.append(f"作者：{author}")
            parts.append("=" * 60)
            parts.append("")

            all_chapters = self.repo.list_chapters()
            all_chapters.sort(key=lambda c: c.get("number", 0))
            for ch_row in all_chapters:
                parts.append(self.export_chapter_plain(ch_row["number"]))
                parts.append("")
                parts.append("=" * 60)
                parts.append("")

        return "\n".join(parts)

    @staticmethod
    def _safe_filename(title: str) -> str:
        """将章节标题转换为安全的文件名"""
        import re
        safe = re.sub(r'[\\/*?:"<>|]', "", title)
        safe = safe.replace(" ", "_").strip("_")
        return safe[:50] if safe else "untitled"

    def export_novel_files(self, output_dir: str, mode: str = "volume",
                           fmt: str = "plain", structure: str = "full") -> dict:
        """导出正文到磁盘文件，按卷分层组织

        这是面向读者的正文导出功能。与返回字符串的 export_* 不同，
        这个方法将内容实际写入文件系统。

        Args:
            output_dir: 输出根目录（将在其下创建小说名子目录）
            mode: 组织方式
                - "volume": 每卷一个文件夹（推荐，大小适中）
                - "chapter": 所有章节放在一个文件夹下
                - "single": 全书合并为单个文件（总集）
                - "auto": 自动判断（≤30章用chapter，>30章用volume）
            fmt: 导出格式 — plain / markdown / html / docx
            structure: 内部结构（仅 mode="single" 时生效）
                - "full": 完整层级：部→卷→章（默认）
                - "flat": 纯章节排列，无卷/部标题（简洁阅读版）

        Returns:
            {
                "status": "ok" | "partial",
                "output_dir": 实际输出路径,
                "mode": 使用的组织方式,
                "format": 格式,
                "volumes": 卷数,
                "chapters_exported": 成功导出的章节数,
                "total_words": 总字数,
                "files": [文件路径列表],
                "structure": 目录结构文本,
            }
        """
        import os
        import re

        self.load_story()
        out = Path(output_dir)
        out.mkdir(parents=True, exist_ok=True)

        # 获取故事信息
        story_data = self.repo.load_story()
        novel_name = story_data.get("name", "未命名") if story_data else "未命名"
        safe_novel = re.sub(r'[\\/*?:"<>|]', "", novel_name).strip() or "小说导出"

        # 决定组织方式
        all_volumes = self.repo.list_volumes()
        all_chapters = self.repo.list_chapters()
        all_chapters.sort(key=lambda c: c.get("number", 0))

        total_chapters = len(all_chapters)
        if mode == "auto":
            mode = "volume" if total_chapters > 30 else "chapter"

        # 计算总字数
        total_words = sum(ch.get("actual_word_count", 0) for ch in all_chapters)

        ext_map = {"plain": ".txt", "markdown": ".md", "html": ".html", "docx": ".docx"}
        ext = ext_map.get(fmt, ".txt")

        files_written = []
        chapters_exported = 0

        # ---- single 模式：全书合并为单个文件 ----
        if mode == "single":
            novel_dir = out / safe_novel
            novel_dir.mkdir(parents=True, exist_ok=True)
            fpath = novel_dir / f"{safe_novel}_full{ext}"

            if fmt == "docx":
                self.export_full_novel_docx(str(fpath), structure=structure)
            elif structure == "flat":
                # 纯章节文本拼接，跳过卷/部结构
                text_exporters = {
                    "plain": self.export_chapter_plain,
                    "markdown": self.export_chapter_markdown,
                    "html": self.export_chapter_html,
                }
                texp = text_exporters.get(fmt, self.export_chapter_plain)
                parts = []
                if fmt == "markdown":
                    parts.append(f"# 《{novel_name}》\n")
                elif fmt == "html":
                    parts.append(
                        f"<!DOCTYPE html><html><head><meta charset=\"utf-8\">"
                        f"<title>{novel_name}</title>"
                        f"<style>body{{max-width:800px;margin:0 auto;padding:20px;"
                        f"font-family:\"Microsoft YaHei\",sans-serif;line-height:1.8}}"
                        f"h1{{text-align:center}}h2{{color:#333}}</style></head><body>"
                        f"<h1>{novel_name}</h1>"
                    )
                else:
                    parts.append(f"《{novel_name}》")
                    parts.append("=" * 60)
                    parts.append("")
                for ch_row in all_chapters:
                    parts.append(texp(ch_row["number"]))
                    parts.append("")
                    if fmt == "markdown":
                        parts.append("---")
                        parts.append("")
                    elif fmt == "html":
                        parts.append("<hr>")
                    else:
                        parts.append("=" * 60)
                        parts.append("")
                if fmt == "html":
                    parts.append("</body></html>")
                fpath.write_text("\n".join(parts), encoding="utf-8")
            else:
                content = self.export_full_novel(fmt)
                fpath.write_text(content, encoding="utf-8")

            chapters_exported = total_chapters
            files_written.append(str(fpath))
            structure_lines = [f"{safe_novel}/", f"  └── {fpath.name}"]

            return {
                "status": "ok",
                "output_dir": str(novel_dir),
                "mode": "single",
                "format": fmt,
                "volumes": len(all_volumes),
                "chapters_exported": chapters_exported,
                "total_chapters": total_chapters,
                "total_words": total_words,
                "files": files_written,
                "structure": "\n".join(structure_lines),
            }

        # ---- volume / chapter 模式 ----
        if fmt == "docx":
            chapter_exporter = self.export_chapter_docx

            def write_chapter_docx(fpath, ch_num):
                """导出单章为 docx 文件"""
                ch_doc = chapter_exporter(ch_num)
                ch_doc.save(str(fpath))
                return True
        else:
            text_exporters = {
                "plain": self.export_chapter_plain,
                "markdown": self.export_chapter_markdown,
                "html": self.export_chapter_html,
            }
            text_exporter = text_exporters.get(fmt, self.export_chapter_plain)

            def write_chapter_text(fpath, ch_num):
                """导出单章为文本文件"""
                content = text_exporter(ch_num)
                if content and content.strip():
                    fpath.write_text(content, encoding="utf-8")
                    return True
                return False

            write_func = write_chapter_text
        if fmt == "docx":
            write_func = write_chapter_docx

        if mode == "volume" and all_volumes:
            # 按卷分层：novel_dir/第N卷 卷名/ch_XXX 章名.ext
            novel_dir = out / safe_novel
            novel_dir.mkdir(parents=True, exist_ok=True)

            vol_chapter_map = {}  # volume_id -> [chapter_rows]

            for v_row in all_volumes:
                vid = v_row["id"]
                ch_rows = self.repo.list_chapters(volume_id=vid)
                ch_rows.sort(key=lambda c: c.get("number", 0))
                vol_chapter_map[vid] = ch_rows

            # 找出无归属章节
            vol_ch_ids = set()
            for ch_rows in vol_chapter_map.values():
                for c in ch_rows:
                    vol_ch_ids.add(c["id"])
            orphan = [c for c in all_chapters if c["id"] not in vol_ch_ids]

            for vol_row in all_volumes:
                vid = vol_row["id"]
                vol_name = f"第{vol_row['number']}卷_{vol_row['title']}"
                safe_vol = re.sub(r'[\\/*?:"<>|]', "", vol_name).strip("_")[:60]
                vol_dir = novel_dir / safe_vol
                vol_dir.mkdir(parents=True, exist_ok=True)

                ch_rows = vol_chapter_map.get(vid, [])
                for ch_row in ch_rows:
                    ch_num = ch_row.get("number", 0)
                    ch_title = ch_row.get("title", "")
                    safe_ch = re.sub(r'[\\/*?:"<>|]', "", f"第{ch_num:02d}章_{ch_title}")
                    fname = f"{safe_ch[:60]}{ext}"
                    fpath = vol_dir / fname

                    if write_func(fpath, ch_num):
                        files_written.append(str(fpath))
                        chapters_exported += 1

            # 无归属章节
            if orphan:
                misc_dir = novel_dir / "独立章节"
                misc_dir.mkdir(parents=True, exist_ok=True)
                for ch_row in orphan:
                    ch_num = ch_row.get("number", 0)
                    ch_title = ch_row.get("title", "")
                    safe_ch = re.sub(r'[\\/*?:"<>|]', "", f"第{ch_num:02d}章_{ch_title}")
                    fname = f"{safe_ch[:60]}{ext}"
                    fpath = misc_dir / fname
                    if write_func(fpath, ch_num):
                        files_written.append(str(fpath))
                        chapters_exported += 1

        else:
            # 单层模式（chapter）：所有章节在一个文件夹下
            novel_dir = out / safe_novel
            novel_dir.mkdir(parents=True, exist_ok=True)

            for ch_row in all_chapters:
                ch_num = ch_row.get("number", 0)
                ch_title = ch_row.get("title", "")
                safe_ch = re.sub(r'[\\/*?:"<>|]', "", f"第{ch_num:02d}章_{ch_title}")
                fname = f"{safe_ch[:60]}{ext}"
                fpath = novel_dir / fname

                if write_func(fpath, ch_num):
                    files_written.append(str(fpath))
                    chapters_exported += 1

        # 生成目录结构描述
        structure_lines = [f"{safe_novel}/"]
        if mode == "volume" and all_volumes:
            for vol_row in all_volumes:
                vid = vol_row["id"]
                vol_name = f"第{vol_row['number']}卷_{vol_row['title']}"
                safe_vol = re.sub(r'[\\/*?:"<>|]', "", vol_name).strip("_")[:60]
                ch_rows = vol_chapter_map.get(vid, [])
                structure_lines.append(f"├── {safe_vol}/")
                for i, ch_row in enumerate(ch_rows):
                    prefix = "│   └──" if i == len(ch_rows) - 1 and not orphan else "│   ├──"
                    structure_lines.append(
                        f"{prefix} 第{ch_row.get('number', 0):02d}章 "
                        f"{ch_row.get('title', '')}{ext}"
                    )
            if orphan:
                structure_lines.append(f"└── 独立章节/")
                for i, ch_row in enumerate(orphan):
                    prefix = "    └──" if i == len(orphan) - 1 else "    ├──"
                    structure_lines.append(
                        f"{prefix} 第{ch_row.get('number', 0):02d}章 "
                        f"{ch_row.get('title', '')}{ext}"
                    )
        else:
            for i, ch_row in enumerate(all_chapters):
                prefix = "└──" if i == len(all_chapters) - 1 else "├──"
                structure_lines.append(
                    f"{prefix} 第{ch_row.get('number', 0):02d}章 "
                    f"{ch_row.get('title', '')}{ext}"
                )

        return {
            "status": "ok" if chapters_exported == total_chapters else "partial",
            "output_dir": str(novel_dir),
            "mode": mode,
            "format": fmt,
            "volumes": len(all_volumes),
            "chapters_exported": chapters_exported,
            "total_chapters": total_chapters,
            "total_words": total_words,
            "files": files_written,
            "structure": "\n".join(structure_lines),
        }

    # ============================================================
    #  DOCX 导出
    # ============================================================

    @staticmethod
    def _docx_styled() -> 'Document':
        """创建一个已配置中文样式的 docx Document"""
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # 默认样式
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Microsoft YaHei'
        font.size = Pt(12)
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.first_line_indent = Cm(0.74)  # 两字符缩进

        # 页边距
        for section in doc.sections:
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(3.18)
            section.right_margin = Cm(3.18)

        return doc

    def export_chapter_docx(self, chapter_number: int, include_summary: bool = False) -> 'Document':
        """导出单章为 python-docx Document 对象

        Args:
            chapter_number: 章节号
            include_summary: 是否包含章节摘要（默认 False，读者版不需要管理元数据）
        """
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        ch = self.get_chapter(chapter_number)
        doc = self._docx_styled()

        if not ch:
            doc.add_paragraph(f"（第{chapter_number}章不存在）")
            return doc

        # 章节标题
        title = doc.add_heading(f'第{ch.number}章  {ch.title}', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 摘要（仅在明确要求时输出，读者版不需要）
        if include_summary and ch.summary:
            quote = doc.add_paragraph(ch.summary)
            quote.style = doc.styles['Quote']
            quote.paragraph_format.first_line_indent = Pt(0)

        # 场景内容
        has_content = any(sc.content.strip() for sc in ch.scenes)
        if has_content:
            scenes_with_content = [sc for sc in ch.scenes if sc.content.strip()]
            multi = len(scenes_with_content) > 1
            first = True
            for sc in scenes_with_content:
                if not first and multi:
                    # 场景分隔线
                    sep = doc.add_paragraph('—' * 20)
                    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    sep.paragraph_format.first_line_indent = Pt(0)
                first = False
                for para_text in sc.content.split('\n'):
                    para_text = para_text.strip()
                    if para_text:
                        doc.add_paragraph(para_text)
        else:
            p = doc.add_paragraph('（本章暂无正文）')
            p.paragraph_format.first_line_indent = Pt(0)

        return doc

    def export_volume_docx(self, volume_number: int) -> 'Document':
        """导出整卷为 python-docx Document 对象（含卷标题 + 全部章节）"""
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        vol_row = self.repo.get_volume_by_number(volume_number)
        doc = self._docx_styled()

        # 卷标题
        if vol_row:
            vol_title = doc.add_heading(
                f'第{vol_row["number"]}卷：{vol_row["title"]}', level=0
            )
            vol_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if vol_row.get("theme"):
                theme_p = doc.add_paragraph(f'主题：{vol_row["theme"]}')
                theme_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                theme_p.paragraph_format.first_line_indent = Pt(0)
            if vol_row.get("summary"):
                quote = doc.add_paragraph(vol_row["summary"])
                quote.style = doc.styles['Quote']
                quote.paragraph_format.first_line_indent = Pt(0)

            chapters = self.repo.list_chapters(volume_id=vol_row["id"])
            chapters.sort(key=lambda c: c.get("number", 0))
        else:
            doc.add_paragraph(f"（第{volume_number}卷不存在）")
            return doc

        if not chapters:
            p = doc.add_paragraph('（本卷暂无章节）')
            p.paragraph_format.first_line_indent = Pt(0)
            return doc

        for ch_row in chapters:
            ch_doc = self.export_chapter_docx(ch_row["number"])
            # 将章节文档的元素追加到卷文档
            self._append_docx_content(doc, ch_doc)
            # 章节间分页
            doc.add_page_break()

        return doc

    def export_full_novel_docx(self, output_path: str, structure: str = "full") -> str:
        """导出全书为单个 DOCX 文件

        Args:
            output_path: 输出文件路径（含 .docx 扩展名）
            structure: 内部结构
                - "full": 部→卷→章 完整层级（默认，适合有复杂结构的作品）
                - "flat": 纯章节排列，无卷/部标题（简洁，适合直接阅读）

        Returns:
            输出文件路径
        """
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        self.load_story()
        story_data = self.repo.load_story()
        title = story_data.get("name", "未命名") if story_data else "未命名"
        author = story_data.get("author", "") if story_data else ""

        doc = self._docx_styled()

        # === 封面 / 书名 ===
        book_title = doc.add_heading(f'《{title}》', level=0)
        book_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if author:
            author_p = doc.add_paragraph(f'作者：{author}')
            author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            author_p.paragraph_format.first_line_indent = Pt(0)

        all_chapters = self.repo.list_chapters()
        all_chapters.sort(key=lambda c: c.get("number", 0))

        if structure == "flat":
            # 纯章节模式：跳过卷/部结构，直接按章节号排列
            for ch_row in all_chapters:
                ch_doc = self.export_chapter_docx(ch_row["number"])
                self._append_docx_content(doc, ch_doc)

        else:
            # === 完整结构：部 → 卷 → 章 ===
            has_parts = len(self.repo.list_parts()) > 0
            has_volumes = len(self.repo.list_volumes()) > 0

            if has_parts:
                for p_row in self.repo.list_parts():
                    # 部标题
                    part_heading = doc.add_heading(
                        f'第{p_row.get("number", 0)}部：{p_row.get("title", "")}', level=1
                    )
                    part_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if p_row.get("summary"):
                        quote = doc.add_paragraph(p_row["summary"])
                        quote.style = doc.styles['Quote']
                        quote.paragraph_format.first_line_indent = Pt(0)

                    for v_row in self.repo.list_volumes(part_id=p_row["id"]):
                        vol_doc = self.export_volume_docx(v_row["number"])
                        self._append_docx_content(doc, vol_doc)

            elif has_volumes:
                for v_row in self.repo.list_volumes():
                    vol_doc = self.export_volume_docx(v_row["number"])
                    self._append_docx_content(doc, vol_doc)

                # 无归属章节
                vol_ch_ids = set()
                for v_row in self.repo.list_volumes():
                    for ch in self.repo.list_chapters(volume_id=v_row["id"]):
                        vol_ch_ids.add(ch["id"])
                orphan = [c for c in all_chapters if c["id"] not in vol_ch_ids]
                if orphan:
                    orphan_heading = doc.add_heading('独立章节', level=1)
                    orphan_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for ch_row in orphan:
                        ch_doc = self.export_chapter_docx(ch_row["number"])
                        self._append_docx_content(doc, ch_doc)
            else:
                # 纯章节模式
                for ch_row in all_chapters:
                    ch_doc = self.export_chapter_docx(ch_row["number"])
                    self._append_docx_content(doc, ch_doc)

        # 保存
        doc.save(output_path)
        return output_path

    @staticmethod
    def _append_docx_content(target: 'Document', source: 'Document') -> None:
        """将 source 文档的所有正文元素追加到 target 文档末尾

        只追加段落和表格等正文元素，跳过标题（避免 docx 内部样式冲突）。
        """
        from docx.oxml.ns import qn

        for element in source.element.body:
            # 跳过第一节的 sectPr（页面设置），让 target 保持自己的页面设置
            if element.tag == qn('w:sectPr'):
                continue
            target.element.body.append(element)
